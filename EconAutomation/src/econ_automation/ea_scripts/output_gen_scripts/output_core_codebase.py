import logging
from pathlib import Path
from typing import Any, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph as ParagraphObject
from docx.table import Table as TableObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Module's logger instance
logger = logging.getLogger(__name__)


def determine_variable_map(
    main_reformatted_data_dict: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    """
    Determines the variable maps for the Word document templates.
    """
    variable_map = {}

    try:
        for wb_name, wb_data in main_reformatted_data_dict.items():
            for variable, value in wb_data.items():
                variable_token = f"[[{variable}]]"
                variable_map[variable_token] = "" if value is None else str(value)
    except Exception as e:
        logger.exception(f"Error determining variable maps: {e}")
        raise

    return variable_map


def add_page_break_before_paragraph(paragraph) -> None:
    """
    Adds a page break before the provided paragraph.
    """
    pPr = (
        paragraph._p.get_or_add_pPr()
    )  # gets/creates the paragraph properties XML element
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreakBefore)


def consolidate_paragraph_runs(paragraph) -> None:
    """
    Consolidates all writeable run text in a paragraph into the first run of that paragraph.
    Paragraph text is preserved and is now editable, but mixed formatting is lost.
    """
    current_group = []

    if not paragraph.runs:
        return

    def run_has_image(run):
        """
        Checks if a run contains an image.
        """
        return run._r.find(qn("w:drawing")) is not None

    def consolidate_group(group):
        """
        Consolidates all writeable run text in a group into the first run of that group.
        """
        if len(group) <= 1:
            return
        full_text = "".join(r.text for r in group)
        group[0].text = full_text
        for r in group[1:]:
            r.text = ""

    for run in paragraph.runs:
        if run_has_image(run):
            consolidate_group(current_group)  # flush group before the image
            current_group = []  # reset for runs after the image
        else:
            current_group.append(run)

    consolidate_group(current_group)  # flush final group


def replace_tokens_in_paragraphs(
    paragraphs: Sequence[ParagraphObject],
    variable_map: dict[str, str],
    target_character: str = "[[",
) -> None:
    """
    Replaces all variable tokens in the document's paragraphs with values from variable_map: {"variable_token": "replacement"}.
    """
    for paragraph in paragraphs:
        if target_character not in paragraph.text:
            continue

        consolidate_paragraph_runs(paragraph)

        for run in paragraph.runs:
            if run._r.find(qn("w:drawing")) is not None:
                continue
            for token, value in variable_map.items():
                run.text = run.text.replace(token, value)


def replace_tokens_in_tables(
    tables: Sequence[TableObject],
    variable_map: dict[str, str],
    target_character: str = "[[",
) -> None:
    """
    Replaces all variable tokens in the document's tables with values from variable_map: {"variable_token": "replacement"}.
    """
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_tokens_in_paragraphs(
                    cell.paragraphs, variable_map, target_character
                )


def autofill_relevant_paragraphs(
    document: DocumentObject,
    variable_map: dict[str, str],
    target_character: str = "[[",
) -> None:
    """
    Executes the final autofill process, replacing all variable tokens in the document with values from variable_map: {"variable_token": "replacement"}.
    """
    try:
        replace_tokens_in_paragraphs(
            paragraphs=document.paragraphs,
            variable_map=variable_map,
            target_character=target_character,
        )
        replace_tokens_in_tables(
            tables=document.tables,
            variable_map=variable_map,
            target_character=target_character,
        )

        for section in document.sections:
            replace_tokens_in_paragraphs(
                paragraphs=section.header.paragraphs,
                variable_map=variable_map,
                target_character=target_character,
            )
            replace_tokens_in_tables(
                tables=section.header.tables,
                variable_map=variable_map,
                target_character=target_character,
            )
            replace_tokens_in_paragraphs(
                paragraphs=section.footer.paragraphs,
                variable_map=variable_map,
                target_character=target_character,
            )
            replace_tokens_in_tables(
                tables=section.footer.tables,
                variable_map=variable_map,
                target_character=target_character,
            )

    except:
        logger.exception("WordTemplateProcessor.autofill_relevant_paragraphs() error")
        raise


def save_output_document(
    document: DocumentObject,
    variable_map: dict[str, str],
    report_type: str,
    selected_output_filepaths: list[Path],
) -> None:
    """
    Saves the output document to the specified paths.
    """
    document_to_save = document
    try:
        for output_filepath in selected_output_filepaths:
            base_stem = f"{variable_map['[[claimant_name_last]]']}{variable_map['[[claimant_name_first_initial]]']} - {report_type}"
            final_output_path = output_filepath / f"{base_stem}.docx"

            output_filepath.mkdir(parents=True, exist_ok=True)

            final_filepath_counter = 0
            while final_output_path.exists():
                final_filepath_counter += 1
                final_output_path = (
                    output_filepath / f"{base_stem}({final_filepath_counter}).docx"
                )

            document_to_save.save(str(final_output_path))

            # DEBUG
            logger.info(
                f"WordTemplateProcessor.save_output_document: Saved output document to: {final_output_path}"
            )
    except KeyError:
        logger.exception("WordTemplateProcessor.save_output_document() KeyError")
        raise
    except TypeError:
        logger.exception("WordTemplateProcessor.save_output_document() TypeError")
        raise
    except FileNotFoundError:
        logger.exception(
            "WordTemplateProcessor.save_output_document() FileNotFoundError"
        )
        raise
    except Exception as e:
        logger.exception(f"WordTemplateProcessor.save_output_document() error: {e}")
        raise


def load_template(template_filepath: Path) -> DocumentObject:
    """
    Loads a Word document template.
    """
    return Document(str(template_filepath))


def AutofillWordTemplates(
    main_reformatted_data_dict: dict[str, dict[str, Any]],
    selected_template_filepaths: list[Path],
    selected_output_filepaths: list[Path],
) -> None:
    """
    Autofills all Word document templates.
    """

    # DEFINING VARIABLES
    documents_dict = {}

    # DETERMINING VARIABLE MAPS
    variable_map = determine_variable_map(
        main_reformatted_data_dict=main_reformatted_data_dict
    )

    # LOADING AND APPENDING ALL WORD DOCUMENT INSTANCES TO DOCUMENTS_DICT
    for template_filepath in selected_template_filepaths:
        documents_dict[template_filepath.name] = load_template(
            template_filepath=template_filepath
        )

    # CALLING METHODS
    try:
        for report_type, document in documents_dict.items():
            autofill_relevant_paragraphs(
                document=document,
                variable_map=variable_map,
                target_character="[[",
            )
            save_output_document(
                document=document,
                variable_map=variable_map,
                report_type=report_type,
                selected_output_filepaths=selected_output_filepaths,
            )
    except Exception as e:
        logger.exception(f"Error autofilling Word document templates: {e}")
        raise

    # def create_relevant_tables(self, extracted_tables_dict: dict[str, pd.DataFrame]=None, target_character="||"):
    #     """
    #     Replaces tables in the document based on the provided tables dictionary.
    #     """
    #     try:
    #         for worksheet_name, tables_dict in extracted_tables_dict.items():
    #             for table_name, table in tables_dict.items():
    #                 self.document.add_table(rows=table.shape[0], cols=table.shape[1])
    #                 self.document.tables[-1].style = "Table Grid"
    #                 for i, row in enumerate(table.iterrows()):
    #                     for j, value in enumerate(row[1]):
    #                     self.document.tables[-1].cell(i, j).text = str(value)
    #         logger.info(f"WordTemplateProcessor.create_relevant_tables: Created tables from {self.template_filepath}")
    #         logger.info(f"WordTemplateProcessor.create_relevant_tables: Created tables dict: {tables_dict}")
    #         return tables_dict
    #     except Exception as e:
    #         logger.exception(f"WordTemplateProcessor.create_relevant_tables: Error creating tables: {e}")
    #         return None
