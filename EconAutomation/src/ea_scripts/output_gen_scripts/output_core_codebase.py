from docx import Document
import logging
from pathlib import Path
from typing import Any

import docx
from docx.document import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Module's logger instance
logger = logging.getLogger(__name__)
        

class WordTemplateProcessor:
    """
    A class that handles the autofilling of Word document templates with extracted data.
    """
    def __init__(self, extracted_data_dict_1: dict[str, Any], extracted_data_dict_2: dict[str, Any] | None=None, template_filepaths_dict: dict[str, Path] = {}, output_save_paths_list: list[Path] = []):
        # DEFINING CLASS ATTRIBUTES
        self.template_filepaths_dict = template_filepaths_dict
        self.output_save_paths_list = output_save_paths_list
        self.documents_list = []
                
        # DETERMINING VARIABLE MAPS
        self.variable_map_1, self.variable_map_2 = self.determine_variable_maps(extracted_data_dict_1, extracted_data_dict_2)
        
        # LOADING AND APPENDING ALL WORD DOCUMENT INSTANCES TO SELF.DOCUMENTS_LIST
        for template_filepath in self.template_filepaths_dict.values():
            self.documents_list.append(self.load_template(template_filepath))
        
        # DEBUG
        for k,v in self.variable_map_1.items():
            logger.debug(f"\nWordTemplateProcessor.__init__.variable_map_1: {k}: {v}")
        
        if self.variable_map_2:
            for k,v in self.variable_map_2.items():
                logger.debug(f"\nWordTemplateProcessor.__init__.variable_map_2: {k}: {v}")
        
        # CALLING METHODS
        try:
            for document in self.documents_list:
                self.autofill_relevant_paragraphs(document)
            self.save_output_document()
        except Exception as e:
            logger.exception(f"Error autofilling Word document templates: {e}")
            raise
        
    def determine_variable_maps(self, extracted_data_dict_1: dict[str, Any], extracted_data_dict_2: dict[str, Any] | None=None) -> tuple[dict[str, Any], dict[str, Any] | None]:
        """
        Determines the variable maps for the Word document templates.
        """
        variable_map_1 = {}
        variable_map_2 = {}
        
        try:
            for k,v in extracted_data_dict_1.items():
                variable_token = f"[[{k}]]"
                variable_map_1[variable_token] = "" if v is None else str(v)

            if extracted_data_dict_2:
                for k,v in extracted_data_dict_2.items():
                    variable_token = f"[[{k}]]"
                    variable_map_2[variable_token] = "" if v is None else str(v)
            else:
                variable_map_2 = None
        except Exception as e:
            logger.exception(f"Error determining variable maps: {e}")
            raise
        
        return variable_map_1, variable_map_2
    
    def load_template(self, template_filepath: Path) -> Document:
        """
        Loads a Word document template.
        """
        return docx.Document(str(template_filepath))
    
    def add_page_break_before_paragraph(self, paragraph) -> None:
        """
        Adds a page break before the provided paragraph.
        """
        pPr = paragraph._p.get_or_add_pPr()  # gets/creates the paragraph properties XML element
        pageBreakBefore = OxmlElement('w:pageBreakBefore')
        pPr.append(pageBreakBefore)
    
    def consolidate_paragraph_runs(self, paragraph) -> None:
        """
        Consolidates all writeable run text in a paragraph into the first run of that paragraph. 
        Paragraph text is preserved and is now editable, but mixed formatting is lost.
        """
        current_group = []
        
        if not paragraph.runs:
            return
        
        def run_has_image(run):
            """
            Checks if a run contains an image.
            """
            return run._r.find(qn('w:drawing')) is not None
        
        def consolidate_group(group):
            """
            Consolidates all writeable run text in a group into the first run of that group.
            """
            if len(group) <= 1:
                return
            full_text = "".join(r.text for r in group)
            group[0].text = full_text
            for r in group[1:]:
                r.text = ""
    
        for run in paragraph.runs:
            if run_has_image(run):
                consolidate_group(current_group)  # flush group before the image
                current_group = []                # reset for runs after the image
            else:
                current_group.append(run)
    
        consolidate_group(current_group)          # flush final group
    
    def replace_tokens_in_paragraphs(self, paragraphs, target_character="[[") -> None:
        """
        Replaces all variable tokens in the document's paragraphs with values from self.variable_map: {"variable_token": "replacement"}.
        """
        for paragraph in paragraphs:
            if target_character not in paragraph.text:
                continue
            
            self.consolidate_paragraph_runs(paragraph)
        
            for run in paragraph.runs:
                if run._r.find(qn('w:drawing')) is not None:
                    continue
                for token, value in self.variable_map_1.items():
                    run.text = run.text.replace(token, value)
                if self.variable_map_2:
                    for token, value in self.variable_map_2.items():
                        run.text = run.text.replace(token, value)
                    
    def replace_tokens_in_tables(self, tables) -> None:
        """
        Replaces all variable tokens in the document's tables with values from self.variable_map: {"variable_token": "replacement"}.
        """
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self.replace_tokens_in_paragraphs(cell.paragraphs)
        
    def autofill_relevant_paragraphs(self, document: Document, target_character="[[") -> None:
        """
        Executes the final autofill process, replacing all variable tokens in the document with values from variable_map: {"variable_token": "replacement"}.
        """
        try:
            self.replace_tokens_in_paragraphs(document.paragraphs, target_character)
            self.replace_tokens_in_tables(document.tables)

            for section in document.sections:
                self.replace_tokens_in_paragraphs(section.header.paragraphs)
                self.replace_tokens_in_tables(section.header.tables)
                self.replace_tokens_in_paragraphs(section.footer.paragraphs)
                self.replace_tokens_in_tables(section.footer.tables)

        except Exception as e:
            logger.exception(f"WordTemplateProcessor.autofill_relevant_paragraphs() error: {e}")
            
    def save_output_document(self) -> None:
        """
        Saves the output document to the specified path.
        """
        try:
            for i, output_save_path in enumerate(self.output_save_paths_list):
                output_path_object = Path(output_save_path)
                base_filename = f"{self.variable_map_1['[[claimant_name_last]]']}{self.variable_map_1['[[claimant_name_first_initial]]']} - {output_path_object.name}"
                final_output_path = Path(rf"{output_path_object}/{base_filename}.docx")

                if not final_output_path.exists():
                    final_output_path.mkdir(parents=True, exist_ok=True)
                
                final_filepath_counter = 0
                while final_output_path.exists():
                    final_filepath_counter += 1
                    final_output_path = Path(rf"{output_path_object}/{base_filename}({final_filepath_counter}).docx")
                
                self.documents_list[i].save(final_output_path)
                
                # DEBUG
                logger.info(f"WordTemplateProcessor.save_output_document: Saved output document to: {final_output_path}")
        except Exception as e:
            logger.exception(f"WordTemplateProcessor.save_output_document() error: {e}")
            raise
            
    # def create_relevant_tables(self, extracted_tables_dict: dict[str, pd.DataFrame]=None, target_character="||"):
    #     """
    #     Replaces tables in the document based on the provided tables dictionary.
    #     """
    #     try:
    #         for worksheet_name, tables_dict in extracted_tables_dict.items():
    #             for table_name, table in tables_dict.items():
    #                 self.document.add_table(rows=table.shape[0], cols=table.shape[1])
    #                 self.document.tables[-1].style = "Table Grid"
    #                 for i, row in enumerate(table.iterrows()):
    #                     for j, value in enumerate(row[1]):
    #                     self.document.tables[-1].cell(i, j).text = str(value)
    #         logger.info(f"WordTemplateProcessor.create_relevant_tables: Created tables from {self.template_filepath}")
    #         logger.info(f"WordTemplateProcessor.create_relevant_tables: Created tables dict: {tables_dict}")
    #         return tables_dict
    #     except Exception as e:
    #         logger.exception(f"WordTemplateProcessor.create_relevant_tables: Error creating tables: {e}")
    #         return None